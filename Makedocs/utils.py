import os
import re
import uuid
from docx import Document


def replace_placeholders(text, data):
    """Заменяет плейсхолдеры в тексте на данные"""
    if not text:
        return text

    for key, value in data.items():
        placeholder = '{{ ' + key + ' }}'
        if placeholder in text:
            text = text.replace(placeholder, str(value or ''))
    return text


def normalize_url(url):
    """Нормализует URL, добавляя протокол если необходимо"""
    if not url:
        return url

    url = url.strip()
    if url and not url.startswith(('http://', 'https://')):
        return f'https://{url}'
    return url


def clean_filename(filename):
    """Очищает имя файла от недопустимых символов, максимально сохраняя читаемость"""
    if not filename:
        return 'document'

    # Убираем расширение, если есть
    name_without_ext = os.path.splitext(filename)[0] if '.' in filename else filename

    # Если исходное имя очень короткое, сохраняем его максимально
    original_name = name_without_ext.strip()
    if original_name and len(original_name) <= 3:
        # Для коротких имён (цифры, короткие слова) - только убираем недопустимые символы
        clean_name = re.sub(r'[<>:"/\\|?*]', '_', original_name)
        if clean_name:
            return clean_name

    # Для длинных имён - полная обработка
    # Заменяем недопустимые символы на подчёркивания
    clean_name = re.sub(r'[<>:"/\\|?*]', '_', name_without_ext)

    # Заменяем последовательности пробелов и подчёркиваний на одно подчёркивание
    clean_name = re.sub(r'[_\s]+', '_', clean_name)

    # Убираем подчёркивания в начале и конце
    clean_name = clean_name.strip('_')

    # Ограничиваем длину (Windows имеет ограничение ~260 символов на полный путь)
    if len(clean_name) > 100:
        clean_name = clean_name[:100].rstrip('_')

    # Возвращаем результат или fallback
    return clean_name if clean_name else 'document'


def create_output_filename_with_client_prefix(client_name, original_filename, templates_count=1, template_name=None):
    """
    Создаёт имя выходного файла, максимально сохраняя оригинальное имя

    Args:
        client_name: Имя клиента
        original_filename: Оригинальное имя файла шаблона
        templates_count: Количество шаблонов (для определения нужности префикса)
        template_name: Имя шаблона (используется при множественных шаблонах)

    Returns:
        str: Имя выходного файла
    """

    # Очищаем имя клиента
    clean_client = clean_filename(client_name)
    if not clean_client:
        clean_client = "Client"

    # Разбираем оригинальное имя файла
    if '.' in original_filename:
        name_part, extension = os.path.splitext(original_filename)
    else:
        name_part, extension = original_filename, '.docx'

    # Очищаем оригинальное имя (но сохраняем его структуру)
    clean_original = clean_filename(name_part)
    if not clean_original:
        clean_original = "document"

    # Логика формирования финального имени:

    # 1. Если шаблон только один - просто добавляем префикс клиента
    if templates_count == 1:
        # Проверяем, есть ли уже имя клиента в имени файла
        if clean_client.lower() in clean_original.lower():
            # Если имя клиента уже есть в файле, оставляем как есть
            final_name = clean_original
        else:
            # Добавляем имя клиента как префикс
            final_name = f"{clean_client}_{clean_original}"

    # 2. Если шаблонов несколько - добавляем и клиента, и шаблон для группировки
    else:
        parts = [clean_client]

        # Добавляем имя шаблона для группировки, если указано
        if template_name:
            clean_template = clean_filename(template_name)
            if clean_template:
                parts.append(clean_template)

        # Добавляем оригинальное имя файла
        parts.append(clean_original)

        final_name = "_".join(parts)

    # Финальная проверка длины (Windows имеет ограничение ~260 символов на полный путь)
    max_name_length = 150  # Оставляем место для расширения и пути
    if len(final_name) > max_name_length:
        # Укорачиваем, сохраняя окончание (которое обычно более важно)
        final_name = "..." + final_name[-(max_name_length - 3):]
        final_name = final_name.lstrip('_')  # Убираем начальные подчёркивания после обрезки

    return f"{final_name}{extension}"


def preserve_original_filename_with_prefix(original_filename, client_name, template_name=None, add_template=False):
    """
    Альтернативная функция - полностью сохраняет оригинальное имя, добавляя только минимальный префикс

    Args:
        original_filename: Оригинальное имя файла
        client_name: Имя клиента
        template_name: Имя шаблона (опционально)
        add_template: Добавлять ли имя шаблона

    Returns:
        str: Имя файла с минимальными изменениями
    """

    # Очищаем имя клиента
    clean_client = clean_filename(client_name)
    if not clean_client:
        clean_client = "Client"

    # Разбираем оригинальное имя
    if '.' in original_filename:
        name_part, extension = os.path.splitext(original_filename)
    else:
        name_part, extension = original_filename, '.docx'

    # Проверяем, нет ли уже имени клиента в файле
    if clean_client.lower() in name_part.lower():
        # Если имя клиента уже есть, возвращаем оригинальное имя
        return original_filename if original_filename.endswith('.docx') else original_filename + '.docx'

    # Формируем префикс
    prefix_parts = [clean_client]

    if add_template and template_name:
        clean_template = clean_filename(template_name)
        if clean_template:
            prefix_parts.append(clean_template)

    prefix = "_".join(prefix_parts)

    # Возвращаем файл с префиксом
    return f"{prefix}_{name_part}{extension}"


def smart_filename_generation(client_name, original_filename, template_name=None, templates_count=1):
    """
    Умная генерация имени файла с учётом контекста

    Логика:
    - Сохраняет оригинальное имя файла максимально
    - Добавляет префикс клиента только если необходимо
    - Учитывает количество шаблонов и файлов
    """

    clean_client = clean_filename(client_name)
    if not clean_client:
        clean_client = "Client"

    # Получаем базовое имя без расширения
    base_name = os.path.splitext(original_filename)[0] if '.' in original_filename else original_filename
    extension = '.docx'

    # Очищаем базовое имя, но сохраняем его читаемость
    clean_base = clean_filename(base_name)
    if not clean_base:
        clean_base = "document"

    # Определяем стратегию именования
    if templates_count == 1:
        # Один шаблон - минимальные изменения
        if len(clean_base) > 3 and not clean_client.lower() in clean_base.lower():
            # Добавляем клиента только если его нет в имени и имя файла достаточно длинное
            result = f"{clean_client}_{clean_base}"
        else:
            # Сохраняем оригинальное имя
            result = clean_base
    else:
        # Несколько шаблонов - нужна группировка
        parts = [clean_client]

        if template_name and len(template_name.strip()) > 0:
            clean_template = clean_filename(template_name)
            if clean_template and len(clean_template) <= 30:  # Ограничиваем длину имени шаблона
                parts.append(clean_template)

        parts.append(clean_base)
        result = "_".join(parts)

    # Контролируем общую длину
    if len(result) > 100:
        # Укорачиваем разумно
        if templates_count == 1:
            # Для одного шаблона - укорачиваем только базовое имя
            max_base_len = 100 - len(clean_client) - 1  # -1 для подчёркивания
            if max_base_len > 10:
                short_base = clean_base[:max_base_len]
                result = f"{clean_client}_{short_base}"
            else:
                result = f"{clean_client}_{clean_base[:50]}"
        else:
            # Для нескольких шаблонов - укорачиваем пропорционально
            result = result[:100]

    return f"{result}{extension}"


def create_simple_output_filename(client_name, original_filename, template_name=None, include_template=False):
    """Создаёт простое и понятное имя для выходного файла"""

    # Очищаем имя клиента
    clean_client = clean_filename(client_name)

    # Очищаем оригинальное имя файла
    if '.' in original_filename:
        name_without_ext = os.path.splitext(original_filename)[0]
    else:
        name_without_ext = original_filename

    clean_original = clean_filename(name_without_ext)

    # Проверяем, что имена не пустые
    if not clean_client:
        clean_client = "Client"
    if not clean_original:
        clean_original = "Document"

    # Формируем финальное имя
    if include_template and template_name:
        clean_template = clean_filename(template_name)
        if clean_template:
            filename_parts = [clean_client, clean_template, clean_original]
        else:
            filename_parts = [clean_client, clean_original]
    else:
        filename_parts = [clean_client, clean_original]

    base_filename = "_".join(filename_parts)

    # Финальная проверка длины
    if len(base_filename) > 150:
        base_filename = base_filename[:150].rstrip('_')

    return f"{base_filename}.docx"


def should_include_template_name(templates_count, files_in_template_count):
    """Определяет, нужно ли включать имя шаблона в имя файла"""

    # Если шаблон только один и файл в нём тоже один - имя шаблона не нужно
    if templates_count == 1 and files_in_template_count == 1:
        return False

    # Если шаблонов несколько - всегда включаем имя для группировки
    if templates_count > 1:
        return True

    # Если шаблон один, но файлов в нём много - включаем имя шаблона для ясности
    if templates_count == 1 and files_in_template_count > 1:
        return True

    return False


def ensure_unique_filename(desired_filename, used_filenames):
    """Обеспечивает уникальность имени файла, добавляя номер при необходимости"""

    if desired_filename not in used_filenames:
        return desired_filename

    # Разбираем имя и расширение
    name_part, ext_part = os.path.splitext(desired_filename)

    # Ищем свободное имя с номером
    counter = 1
    while counter <= 999:  # Разумное ограничение
        numbered_filename = f"{name_part}_{counter:02d}{ext_part}"
        if numbered_filename not in used_filenames:
            return numbered_filename
        counter += 1

    # Если дошли до сюда, используем UUID
    unique_id = str(uuid.uuid4())[:8]
    return f"{name_part}_{unique_id}{ext_part}"


def advanced_replace_in_paragraph(paragraph, data):
    """
    Продвинутая замена плейсхолдеров в параграфе, которая правильно обрабатывает
    разбитые на несколько run'ов плейсхолдеры
    """
    if not paragraph or not hasattr(paragraph, 'runs'):
        return False

    if not paragraph.runs:
        return False

    # Получаем весь текст параграфа
    full_text = paragraph.text
    if not full_text or '{{' not in full_text:
        return False

    # Заменяем плейсхолдеры в тексте
    original_text = full_text
    replacements_made = 0

    for key, value in data.items():
        placeholder = '{{ ' + key + ' }}'
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value or ''))
            replacements_made += 1

    # Если изменений не было
    if full_text == original_text:
        return False

    # Теперь обновляем параграф
    try:
        # Очищаем все существующие run'ы
        for run in paragraph.runs:
            run.text = ""

        # Добавляем новый текст в первый run, сохраняя его форматирование
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)

        return True

    except Exception as e:
        return False


def process_table_with_merged_cells(table, data):
    """
    Улучшенная обработка таблиц с учётом объединённых ячеек
    """
    replacements_made = 0

    try:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                try:
                    # Получаем весь текст ячейки
                    cell_text = cell.text

                    if '{{' in cell_text:
                        # Обрабатываем все параграфы в ячейке
                        cell_replacements = 0
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if advanced_replace_in_paragraph(paragraph, data):
                                cell_replacements += 1
                                replacements_made += 1

                        if cell_replacements == 0:
                            # Попробуем альтернативный метод для проблемных ячеек
                            if try_alternative_cell_replacement(cell, data):
                                replacements_made += 1

                except Exception as e:
                    continue

    except Exception as e:
        pass

    return replacements_made


def try_alternative_cell_replacement(cell, data):
    """
    Альтернативный метод замены для проблемных ячеек
    """
    try:
        # Получаем весь текст ячейки
        cell_text = cell.text

        if not cell_text or '{{' not in cell_text:
            return False

        # Выполняем замены
        modified_text = cell_text
        for key, value in data.items():
            placeholder = '{{ ' + key + ' }}'
            if placeholder in modified_text:
                modified_text = modified_text.replace(placeholder, str(value or ''))

        # Если изменений не было
        if modified_text == cell_text:
            return False

        # Очищаем ячейку и добавляем новый текст
        cell.clear()
        paragraph = cell.paragraphs[0]  # После clear() остается один пустой параграф
        paragraph.add_run(modified_text)

        return True

    except Exception as e:
        return False


def enhanced_process_docx_template(template_path, output_path, data):
    """
    Улучшенная версия обработки DOCX с продвинутой заменой плейсхолдеров
    """
    try:
        doc = Document(template_path)

        replacements_made = 0

        # Обработка основного текста с улучшенным алгоритмом
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                if advanced_replace_in_paragraph(paragraph, data):
                    replacements_made += 1
            except Exception as e:
                continue

        # Улучшенная обработка таблиц
        for table_idx, table in enumerate(doc.tables):
            table_replacements = process_table_with_merged_cells(table, data)
            replacements_made += table_replacements

        # Обработка заголовков и футеров
        for section_idx, section in enumerate(doc.sections):
            try:
                # Заголовки
                if hasattr(section, 'header') and section.header:
                    for para_idx, paragraph in enumerate(section.header.paragraphs):
                        try:
                            if advanced_replace_in_paragraph(paragraph, data):
                                replacements_made += 1
                        except Exception as e:
                            continue

                # Футеры
                if hasattr(section, 'footer') and section.footer:
                    for para_idx, paragraph in enumerate(section.footer.paragraphs):
                        try:
                            if advanced_replace_in_paragraph(paragraph, data):
                                replacements_made += 1
                        except Exception as e:
                            continue
            except Exception as e:
                continue

        # Сохраняем документ
        doc.save(output_path)

        return True

    except Exception as e:
        raise e


def process_docx_template_safe(template_path, output_path, data):
    """
    Обертка для обратной совместимости - использует улучшенную версию
    """
    return enhanced_process_docx_template(template_path, output_path, data)